from __future__ import annotations

import argparse
import csv
import math
import re
import sys
from dataclasses import dataclass
from pathlib import Path
from zipfile import ZipFile

import matplotlib

matplotlib.use("Agg")
import matplotlib.pyplot as plt

PROJECT_ROOT = Path(__file__).resolve().parents[1]
if str(PROJECT_ROOT) not in sys.path:
    sys.path.insert(0, str(PROJECT_ROOT))

LOG_DIR = PROJECT_ROOT / "logs" / "strapdown_accuracy"
REPORT_PATH = PROJECT_ROOT / "完整方案" / "MNP_SITL固定Vm比例导引识别对比报告.md"
DOCX_PATH = PROJECT_ROOT / "完整方案" / "MNP_SITL固定Vm比例导引识别对比报告.docx"
ASSET_DIR = PROJECT_ROOT / "完整方案" / "assets" / "MNP_SITL固定Vm比例导引识别对比报告"
GRAVITY_MPS2 = 9.80665

CASE_LABELS = {
    "M": "M sensor-noise KF",
    "N": "N sensor-noise raw",
    "P": "P no-sensor-noise raw",
}


@dataclass
class CaseRow:
    label: str
    case: str
    start_range_m: float
    hit: bool
    hit_t_s: float
    min_range_m: float
    final_range_m: float
    frames: int
    detected_frames: int
    valid_frames: int
    avg_sim_sample_fps: float
    max_load_factor_fd_g: float
    csv_path: Path


def _float(value: object, default: float = math.nan) -> float:
    if value is None or value == "":
        return default
    try:
        result = float(value)
    except (TypeError, ValueError):
        return default
    return result if math.isfinite(result) else default


def _int(value: object, default: int = 0) -> int:
    number = _float(value)
    return default if not math.isfinite(number) else int(number)


def _bool(value: object) -> bool:
    return str(value or "").strip().lower() in {"1", "true", "yes"}


def _read_csv(path: Path) -> list[dict[str, str]]:
    with path.open(newline="", encoding="utf-8") as stream:
        return list(csv.DictReader(stream))


def _resolve(path_text: str) -> Path:
    path = Path(path_text)
    return path if path.is_absolute() else PROJECT_ROOT / path


def _rel(path: Path) -> str:
    return path.relative_to(REPORT_PATH.parent).as_posix()


def _series(rows: list[dict[str, str]], key: str) -> list[float]:
    return [_float(row.get(key)) for row in rows]


def _finite(values: list[float]) -> list[float]:
    return [value for value in values if math.isfinite(value)]


def _percentile(values: list[float], q: float) -> float:
    finite = sorted(_finite(values))
    if not finite:
        return math.nan
    if len(finite) == 1:
        return finite[0]
    position = (len(finite) - 1) * max(0.0, min(1.0, q))
    lower = int(math.floor(position))
    upper = int(math.ceil(position))
    if lower == upper:
        return finite[lower]
    fraction = position - lower
    return finite[lower] * (1.0 - fraction) + finite[upper] * fraction


def _vector(row: dict[str, str], keys: tuple[str, str, str]) -> tuple[float, float, float] | None:
    values = tuple(_float(row.get(key)) for key in keys)
    if not all(math.isfinite(value) for value in values):
        return None
    return values


def _norm_series(rows: list[dict[str, str]], keys: tuple[str, str, str]) -> list[float]:
    result: list[float] = []
    for row in rows:
        vec = _vector(row, keys)
        if vec is not None:
            result.append(math.sqrt(sum(value * value for value in vec)))
    return result


def _normalize(vec: tuple[float, float, float]) -> tuple[float, float, float] | None:
    norm = math.sqrt(sum(value * value for value in vec))
    if norm <= 1.0e-9:
        return None
    return tuple(value / norm for value in vec)


def _los_sep_deg(a: tuple[float, float, float] | None, b: tuple[float, float, float] | None) -> float:
    if a is None or b is None:
        return math.nan
    an = _normalize(a)
    bn = _normalize(b)
    if an is None or bn is None:
        return math.nan
    dot = max(-1.0, min(1.0, sum(an[index] * bn[index] for index in range(3))))
    return math.degrees(math.acos(dot))


def _los_errors(rows: list[dict[str, str]]) -> list[float]:
    errors: list[float] = []
    for row in rows:
        visual = _vector(row, ("lambda_x", "lambda_y", "lambda_z"))
        camera = _vector(row, ("camera_world_x", "camera_world_y", "camera_world_z"))
        intruder = _vector(row, ("intruder_x", "intruder_y", "intruder_z"))
        truth = None
        if camera is not None and intruder is not None:
            truth = tuple(intruder[index] - camera[index] for index in range(3))
        errors.append(_los_sep_deg(visual, truth))
    return errors


def _required_load(rows: list[dict[str, str]]) -> list[float]:
    result: list[float] = []
    prev_t: float | None = None
    prev_v: tuple[float, float, float] | None = None
    for row in rows:
        t = _float(row.get("t"))
        v = _vector(row, ("v_cmd_x", "v_cmd_y", "v_cmd_z"))
        if not math.isfinite(t) or v is None or prev_t is None or prev_v is None:
            result.append(0.0 if math.isfinite(t) and v is not None else math.nan)
        else:
            dt = t - prev_t
            if dt <= 1.0e-6:
                result.append(math.nan)
            else:
                dv = math.sqrt(sum((v[index] - prev_v[index]) ** 2 for index in range(3)))
                result.append(dv / dt / GRAVITY_MPS2)
        if math.isfinite(t) and v is not None:
            prev_t = t
            prev_v = v
    return result


def _g_eval_load(rows: list[dict[str, str]]) -> list[float]:
    loads: list[float] = []
    for row in rows:
        vec = _vector(row, ("g_eval_x", "g_eval_y", "g_eval_z"))
        if vec is None:
            loads.append(math.nan)
        else:
            loads.append(math.sqrt(sum(value * value for value in vec)) / GRAVITY_MPS2)
    return loads


def _load_summary(path: Path, label: str) -> list[CaseRow]:
    if not path.exists():
        return []
    rows: list[CaseRow] = []
    for item in _read_csv(path):
        rows.append(
            CaseRow(
                label=label,
                case=item.get("case", ""),
                start_range_m=_float(item.get("start_horizontal_range_m")),
                hit=_bool(item.get("hit")),
                hit_t_s=_float(item.get("hit_t_s")),
                min_range_m=_float(item.get("min_range_m")),
                final_range_m=_float(item.get("final_range_m")),
                frames=_int(item.get("frames")),
                detected_frames=_int(item.get("detected_frames")),
                valid_frames=_int(item.get("valid_frames")),
                avg_sim_sample_fps=_float(item.get("avg_sim_sample_fps")),
                max_load_factor_fd_g=_float(item.get("max_load_factor_fd_g")),
                csv_path=_resolve(item.get("csv_path", "")),
            )
        )
    return rows


def load_rows(stamp: str) -> list[CaseRow]:
    rows: list[CaseRow] = []
    for label in ("M", "N", "P"):
        rows.extend(_load_summary(LOG_DIR / f"strapdown_clock1_sitl_{label}_{stamp}_summary.csv", label))
    return sorted(rows, key=lambda row: (row.start_range_m, row.label))


def _first_sample(rows: list[CaseRow]) -> dict[str, str]:
    for row in rows:
        if row.csv_path.exists():
            samples = _read_csv(row.csv_path)
            if samples:
                return samples[0]
    return {}


def _summary_table(rows: list[CaseRow]) -> str:
    lines = [
        "|工况|命中数|命中距离m|未命中距离m|最小中心距离m|有效帧/总帧|检测帧/总帧|",
        "|---|---:|---|---|---:|---:|---:|",
    ]
    for label in ("M", "N", "P"):
        group = [row for row in rows if row.label == label]
        if not group:
            continue
        hit_ranges = ", ".join(f"{row.start_range_m:.0f}" for row in group if row.hit) or "-"
        miss_ranges = ", ".join(f"{row.start_range_m:.0f}" for row in group if not row.hit) or "-"
        min_range = min(row.min_range_m for row in group if math.isfinite(row.min_range_m))
        valid_frames = sum(row.valid_frames for row in group)
        detected_frames = sum(row.detected_frames for row in group)
        frames = sum(row.frames for row in group)
        lines.append(
            f"|{label}|{sum(row.hit for row in group)}/{len(group)}|{hit_ranges}|{miss_ranges}|"
            f"{min_range:.3f}|{valid_frames}/{frames}|{detected_frames}/{frames}|"
        )
    return "\n".join(lines)


def _detail_table(rows: list[CaseRow]) -> str:
    lines = [
        "|工况|距离m|碰撞|碰撞时间s|最小中心距离m|终点距离m|检测帧/总帧|有效帧|实际过载max g|速度指令P95 g|固定Vm理论P95 g|LOS误差P95 deg|sim FPS|",
        "|---|---:|---:|---:|---:|---:|---:|---:|---:|---:|---:|---:|---:|",
    ]
    for row in rows:
        samples = _read_csv(row.csv_path)
        lines.append(
            f"|{row.label}|{row.start_range_m:.0f}|{1 if row.hit else 0}|"
            f"{'-' if not math.isfinite(row.hit_t_s) else f'{row.hit_t_s:.2f}'}|"
            f"{row.min_range_m:.3f}|{row.final_range_m:.3f}|{row.detected_frames}/{row.frames}|{row.valid_frames}|"
            f"{row.max_load_factor_fd_g:.2f}|{_percentile(_required_load(samples), 0.95):.2f}|"
            f"{_percentile(_g_eval_load(samples), 0.95):.2f}|{_percentile(_los_errors(samples), 0.95):.2f}|"
            f"{row.avg_sim_sample_fps:.2f}|"
        )
    return "\n".join(lines)


def _mode_check_markdown(rows: list[CaseRow]) -> str:
    bad_ttc = 0
    total = 0
    laws: dict[str, int] = {}
    modes: dict[str, int] = {}
    for row in rows:
        for sample in _read_csv(row.csv_path):
            total += 1
            law = sample.get("guidance_law", "")
            mode = sample.get("guidance_mode", "")
            laws[law] = laws.get(law, 0) + 1
            modes[mode] = modes.get(mode, 0) + 1
            if mode == "ttc_png" or sample.get("ttc_used_for_guidance") == "1":
                bad_ttc += 1
    law_text = ", ".join(f"`{key}`={value}" for key, value in sorted(laws.items())) or "-"
    mode_text = ", ".join(f"`{key}`={value}" for key, value in sorted(modes.items())) or "-"
    return "\n".join(
        [
            f"- 总帧数 `{total}`，TTC 导引帧数 `{bad_ttc}`。",
            f"- `guidance_law` 分布：{law_text}。",
            f"- `guidance_mode` 分布：{mode_text}。",
        ]
    )


def plot_summary(rows: list[CaseRow], output: Path) -> None:
    fig, axes = plt.subplots(2, 2, figsize=(14, 8.5))
    ax_min, ax_hit, ax_valid, ax_detect = axes.flat
    for label in ("M", "N", "P"):
        group = [row for row in rows if row.label == label]
        if not group:
            continue
        x = [row.start_range_m for row in group]
        ax_min.plot(x, [row.min_range_m for row in group], marker="o", label=CASE_LABELS[label])
        ax_hit.plot(x, [1 if row.hit else 0 for row in group], marker="o", label=CASE_LABELS[label])
        ax_valid.plot(x, [100.0 * row.valid_frames / max(1, row.frames) for row in group], marker="o", label=CASE_LABELS[label])
        ax_detect.plot(x, [100.0 * row.detected_frames / max(1, row.frames) for row in group], marker="o", label=CASE_LABELS[label])
    ax_min.set_title("Minimum center range")
    ax_min.set_ylabel("m")
    ax_hit.set_title("AirSim collision hit")
    ax_hit.set_ylabel("hit")
    ax_valid.set_title("Valid LOS frame ratio")
    ax_valid.set_ylabel("%")
    ax_detect.set_title("Detection frame ratio")
    ax_detect.set_ylabel("%")
    for ax in axes.flat:
        ax.grid(True, alpha=0.3)
        ax.set_xlabel("Initial horizontal range / m")
        ax.legend(fontsize=8)
    fig.tight_layout()
    fig.savefig(output, dpi=170)
    plt.close(fig)


def plot_per_distance(rows: list[CaseRow], output_dir: Path) -> dict[float, Path]:
    images: dict[float, Path] = {}
    for range_m in sorted({row.start_range_m for row in rows if math.isfinite(row.start_range_m)}):
        group = [row for row in rows if row.start_range_m == range_m]
        fig, axes = plt.subplots(5, 1, figsize=(12, 15), sharex=False)
        ax_actual, ax_req, ax_los, ax_range, ax_detect = axes
        for item in group:
            samples = _read_csv(item.csv_path)
            t = _series(samples, "t")
            label = CASE_LABELS[item.label]
            ax_actual.plot(t, _series(samples, "load_factor_fd_g"), linewidth=1.0, label=label)
            ax_req.plot(t, _g_eval_load(samples), linewidth=1.0, label=label)
            ax_los.plot(t, _los_errors(samples), linewidth=1.0, label=label)
            ax_range.plot(t, _series(samples, "range"), linewidth=1.1, label=f"{label} {'hit' if item.hit else 'miss'}")
            ax_detect.plot(t, _series(samples, "bbox_area"), linewidth=1.0, label=label)
        ax_actual.set_title(f"{range_m:.0f}m actual overload")
        ax_actual.set_ylabel("g")
        ax_req.set_title("Fixed Vm PNG theoretical acceleration demand")
        ax_req.set_ylabel("g")
        ax_los.set_title("Visual LOS vs camera-origin truth LOS")
        ax_los.set_ylabel("deg")
        ax_range.set_title("True center range")
        ax_range.set_ylabel("m")
        ax_detect.set_title("BBox area ratio")
        ax_detect.set_xlabel("Time / s")
        ax_detect.set_ylabel("area")
        for ax in axes:
            ax.grid(True, alpha=0.3)
            ax.legend(fontsize=5, ncol=2)
        fig.tight_layout()
        path = output_dir / f"mnp_compare_{int(round(range_m)):03d}m.png"
        fig.savefig(path, dpi=170)
        plt.close(fig)
        images[range_m] = path
    return images


def _settings_markdown(rows: list[CaseRow], stamp: str) -> str:
    sample = _first_sample(rows)
    vm = _float(sample.get("speed_cap"), 10.0)
    nav = _float(sample.get("navigation_constant"), 3.0)
    gain = _float(sample.get("vm_png_gain"), vm * nav)
    return "\n".join(
        [
            "|项目|M|N|P|",
            "|---|---|---|---|",
            "|对应 J/K/L|J|K|L|",
            "|LOS Kalman|开启|关闭|关闭|",
            "|IMU/GPS 噪声|开启|开启|关闭|",
            "|导引模式|`fixed_vm_png`|`fixed_vm_png`|`fixed_vm_png`|",
            f"|导航比 N|`{nav:g}`|`{nav:g}`|`{nav:g}`|",
            f"|Vm|`{vm:g} m/s`|`{vm:g} m/s`|`{vm:g} m/s`|",
            f"|固定增益 N*Vm|`{gain:g}`|`{gain:g}`|`{gain:g}`|",
            "|TTC|禁用|禁用|禁用|",
            "|bbox center 噪声|`3.0px`|`3.0px`|`3.0px`|",
            "|bbox area 噪声|`8.0%`|`8.0%`|`8.0%`|",
            "|Actor 名义尺寸|`1m x 1m x 0.5m`|`1m x 1m x 0.5m`|`1m x 1m x 0.5m`|",
            f"|实验 stamp|`{stamp}`|`{stamp}`|`{stamp}`|",
        ]
    )


def _conclusion(rows: list[CaseRow]) -> str:
    lines: list[str] = []
    for label in ("M", "N", "P"):
        group = [row for row in rows if row.label == label]
        if not group:
            continue
        hits = sum(row.hit for row in group)
        hit_ranges = ", ".join(f"{row.start_range_m:.0f}m" for row in group if row.hit) or "-"
        miss_ranges = ", ".join(f"{row.start_range_m:.0f}m" for row in group if not row.hit) or "-"
        samples: list[dict[str, str]] = []
        for row in group:
            samples.extend(_read_csv(row.csv_path))
        valid_ratio = sum(row.valid_frames for row in group) / max(1, sum(row.frames for row in group))
        detect_ratio = sum(row.detected_frames for row in group) / max(1, sum(row.frames for row in group))
        req_p95 = _percentile(_g_eval_load(samples), 0.95)
        los_p95 = _percentile(_los_errors(samples), 0.95)
        lines.append(
            f"- {label} 组命中 `{hits}/{len(group)}`，命中距离 `{hit_ranges}`，未命中距离 `{miss_ranges}`，"
            f"检测帧比例 `{100.0 * detect_ratio:.1f}%`，有效 LOS 帧比例 `{100.0 * valid_ratio:.1f}%`，"
            f"固定 Vm 理论过载 P95 `{req_p95:.2f}g`，LOS 误差 P95 `{los_p95:.2f}deg`。"
        )
    lines.append(
        "- 本轮实验的主要目的不是替代 J/K/L，而是隔离 TTC 面积通道。若 M/N/P 的检测帧比例足够高但命中率下降，问题更可能来自固定 Vm 导引强度、PX4 速度响应或末端 LOS 几何，而不是 TTC 估计。"
    )
    lines.append(
        "- 若 M 明显优于 N，说明 bbox 噪声下 LOS Kalman 对固定 Vm PNG 有帮助；若 P 明显优于 N，说明 IMU/GPS 噪声经 PX4 EKF 和姿态链路影响了捷联视觉导引。"
    )
    return "\n".join(lines)


def write_report(rows: list[CaseRow], stamp: str, summary_img: Path, per_distance: dict[float, Path]) -> None:
    REPORT_PATH.parent.mkdir(parents=True, exist_ok=True)
    image_lines = "\n".join(f"![{int(range_m)}m]({_rel(path)})" for range_m, path in sorted(per_distance.items()))
    report = f"""# M/N/P PX4 SITL 固定 Vm 比例导引识别对比报告

## 1. 实验目的

本报告只比较 M/N/P 三组。三组均使用 PX4 SITL 拦截机、按轨迹移动的 `1m x 1m x 0.5m` Actor 目标、中心捷联相机、bbox 中等噪声、`ClockSpeed=1.0`、初始距离 `40-140m`、高度差 `20m`、目标速度 `5m/s`、拦截机速度上限 `10m/s`。

本轮算法层面不使用 TTC。导引律为固定 `V_m` 型比例导引：`a_req = N * V_m * (omega_LOS x lambda)`，其中 `N=3`、`V_m=10m/s`。在 AirSim/PX4 速度控制实现中，该加速度需求映射为现有速度指令修正，并保留横向/纵向限幅与末端外推状态机。

## 2. 工况设置

{_settings_markdown(rows, stamp)}

## 3. TTC 禁用核查

{_mode_check_markdown(rows)}

## 4. 总览图

![summary]({_rel(summary_img)})

## 5. 命中汇总

{_summary_table(rows)}

## 6. 明细表

{_detail_table(rows)}

## 7. 单距离曲线

每个距离一张图。过载子图包含实际过载和固定 `V_m` PNG 理论加速度需求；报告不绘制 TTC 曲线。

{image_lines}

## 8. 结论

{_conclusion(rows)}
"""
    REPORT_PATH.write_text(report, encoding="utf-8")


def _convert_to_docx() -> None:
    try:
        from docx import Document
        from docx.enum.section import WD_ORIENT
        from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
        from docx.shared import Cm, Pt
    except ImportError:
        print("python-docx not installed; skip docx export")
        return

    text = REPORT_PATH.read_text(encoding="utf-8")
    doc = Document()
    section = doc.sections[0]
    section.orientation = WD_ORIENT.PORTRAIT
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(1.5)
    section.bottom_margin = Cm(1.5)
    section.left_margin = Cm(1.3)
    section.right_margin = Cm(1.3)

    for style_name in ("Normal", "Body Text"):
        if style_name in doc.styles:
            style = doc.styles[style_name]
            style.font.name = "Arial"
            style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
            style.font.size = Pt(10.5)
    for index in range(1, 5):
        style_name = f"Heading {index}"
        if style_name in doc.styles:
            style = doc.styles[style_name]
            style.font.name = "Arial"
            style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
            style.font.bold = True

    heading_re = re.compile(r"^(#{1,6})\s+(.*)$")
    image_re = re.compile(r"^!\[([^\]]*)\]\(([^)]+)\)\s*$")
    token_re = re.compile(r"(`[^`]+`|\*\*[^*]+\*\*)")

    def add_runs(paragraph, value: str) -> None:
        pos = 0
        for match in token_re.finditer(value):
            if match.start() > pos:
                paragraph.add_run(value[pos : match.start()])
            token = match.group(0)
            if token.startswith("`"):
                run = paragraph.add_run(token[1:-1])
                run.font.name = "Consolas"
                run._element.rPr.rFonts.set(qn("w:eastAsia"), "Consolas")
            else:
                run = paragraph.add_run(token[2:-2])
                run.bold = True
            pos = match.end()
        if pos < len(value):
            paragraph.add_run(value[pos:])

    def set_cell_text(cell, value: str, bold: bool = False) -> None:
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(value.replace("`", ""))
        run.bold = bold
        run.font.size = Pt(8)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

    def shade(cell) -> None:
        tc_pr = cell._tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:fill"), "D9EAF7")
        tc_pr.append(shd)

    lines = text.splitlines()
    index = 0
    while index < len(lines):
        line = lines[index].rstrip()
        if not line:
            index += 1
            continue
        heading = heading_re.match(line)
        if heading:
            level = min(len(heading.group(1)), 4)
            doc.add_heading(heading.group(2), level=level)
            index += 1
            continue
        image = image_re.match(line)
        if image:
            label, rel = image.groups()
            path = REPORT_PATH.parent / rel
            if path.exists():
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.add_run().add_picture(str(path), width=Cm(18.0))
                if label:
                    cap = doc.add_paragraph(label)
                    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
            index += 1
            continue
        if line.startswith("|"):
            table_lines = []
            while index < len(lines) and lines[index].strip().startswith("|"):
                table_lines.append(lines[index].strip())
                index += 1
            raw_rows = [[cell.strip() for cell in item.strip("|").split("|")] for item in table_lines]
            if len(raw_rows) >= 2:
                body = [raw_rows[0]] + raw_rows[2:]
                cols = max(len(row) for row in body)
                table = doc.add_table(rows=len(body), cols=cols)
                table.alignment = WD_TABLE_ALIGNMENT.CENTER
                table.style = "Table Grid"
                for row_i, row in enumerate(body):
                    for col_i in range(cols):
                        cell = table.cell(row_i, col_i)
                        set_cell_text(cell, row[col_i] if col_i < len(row) else "", row_i == 0)
                        if row_i == 0:
                            shade(cell)
                doc.add_paragraph()
            continue
        if line.startswith("- "):
            p = doc.add_paragraph(style="List Bullet")
            add_runs(p, line[2:])
            index += 1
            continue
        p = doc.add_paragraph()
        add_runs(p, line)
        index += 1

    doc.save(DOCX_PATH)
    with ZipFile(DOCX_PATH) as archive:
        media = [name for name in archive.namelist() if name.startswith("word/media/")]
    print(f"docx={DOCX_PATH}")
    print(f"docx_media_count={len(media)}")


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description="Generate M/N/P fixed-Vm strapdown comparison report.")
    parser.add_argument("--stamp", required=True)
    parser.add_argument("--no-docx", action="store_true")
    return parser.parse_args()


def main() -> None:
    args = parse_args()
    rows = load_rows(args.stamp)
    if not rows:
        raise SystemExit("no M/N/P rows found")
    ASSET_DIR.mkdir(parents=True, exist_ok=True)
    summary_img = ASSET_DIR / f"mnp_summary_{args.stamp}.png"
    plot_summary(rows, summary_img)
    per_distance = plot_per_distance(rows, ASSET_DIR)
    write_report(rows, args.stamp, summary_img, per_distance)
    print(f"report={REPORT_PATH}")
    print(f"stamp={args.stamp}")
    if not args.no_docx:
        _convert_to_docx()


if __name__ == "__main__":
    main()
