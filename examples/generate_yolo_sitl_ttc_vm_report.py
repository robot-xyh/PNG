from __future__ import annotations

import argparse
import csv
import json
import math
import re
import sys
from dataclasses import dataclass
from pathlib import Path
from zipfile import ZipFile

import matplotlib

matplotlib.use("Agg")
import matplotlib.pyplot as plt
import numpy as np

PROJECT_ROOT = Path(__file__).resolve().parents[1]
if str(PROJECT_ROOT) not in sys.path:
    sys.path.insert(0, str(PROJECT_ROOT))

LOG_DIR = PROJECT_ROOT / "logs" / "yolo_sitl_ttc_vm"
REPORT_PATH = PROJECT_ROOT / "完整方案" / "YOLO_SITL_TTC_VM拦截对比报告.md"
ASSET_DIR = PROJECT_ROOT / "完整方案" / "assets" / "YOLO_SITL_TTC_VM拦截对比报告"
DEFAULT_TITLE = "YOLO + ByteTrack PX4 SITL TTC / V_m 拦截对比报告"
DEFAULT_RANGE_NOTE = "两组均测试 50m、60m、70m、80m、90m、100m，每个工况重启 PX4 SITL 和 Blocks。"
GRAVITY_MPS2 = 9.80665

LABELS = {
    "TTC": "TTC + LOS/Vm soft guidance",
    "VM": "fixed Vm PNG",
}


@dataclass
class CaseRow:
    label: str
    case: str
    start_range_m: float
    hit: bool
    hit_t_s: float
    near_hit: bool
    near_hit_t_s: float
    near_hit_range_m: float
    min_range_m: float
    final_range_m: float
    frames: int
    detected_frames: int
    valid_frames: int
    avg_wall_fps: float
    avg_sim_sample_fps: float
    avg_detector_fps: float
    max_load_factor_fd_g: float
    csv_path: Path
    meta_path: Path


def _float(value: object, default: float = math.nan) -> float:
    if value is None or value == "":
        return default
    try:
        number = float(value)
    except (TypeError, ValueError):
        return default
    return number if math.isfinite(number) else default


def _int(value: object, default: int = 0) -> int:
    number = _float(value)
    return default if not math.isfinite(number) else int(number)


def _bool(value: object) -> bool:
    return str(value or "").strip().lower() in {"1", "true", "yes"}


def _read_csv(path: Path) -> list[dict[str, str]]:
    with path.open(newline="", encoding="utf-8") as stream:
        return list(csv.DictReader(stream))


def _read_json(path: Path) -> dict:
    if not path.exists():
        return {}
    with path.open(encoding="utf-8") as stream:
        return json.load(stream)


def _resolve(path_text: str) -> Path:
    path = Path(path_text)
    return path if path.is_absolute() else PROJECT_ROOT / path


def _rel(path: Path) -> str:
    return path.relative_to(REPORT_PATH.parent).as_posix()


def _series(rows: list[dict[str, str]], key: str) -> list[float]:
    return [_float(row.get(key)) for row in rows]


def _finite(values: list[float]) -> list[float]:
    return [value for value in values if math.isfinite(value)]


def _percentile(values: list[float], q: float) -> float:
    finite = sorted(_finite(values))
    if not finite:
        return math.nan
    if len(finite) == 1:
        return finite[0]
    pos = (len(finite) - 1) * max(0.0, min(1.0, q))
    lo = int(math.floor(pos))
    hi = int(math.ceil(pos))
    if lo == hi:
        return finite[lo]
    frac = pos - lo
    return finite[lo] * (1.0 - frac) + finite[hi] * frac


def _vector(row: dict[str, str], keys: tuple[str, str, str]) -> tuple[float, float, float] | None:
    values = tuple(_float(row.get(key)) for key in keys)
    if not all(math.isfinite(value) for value in values):
        return None
    return values


def _norm_series(rows: list[dict[str, str]], keys: tuple[str, str, str]) -> list[float]:
    result: list[float] = []
    for row in rows:
        vec = _vector(row, keys)
        if vec is not None:
            result.append(math.sqrt(sum(component * component for component in vec)))
    return result


def _required_load(rows: list[dict[str, str]]) -> list[float]:
    result: list[float] = []
    prev_t: float | None = None
    prev_v: tuple[float, float, float] | None = None
    for row in rows:
        t = _float(row.get("t"))
        v = _vector(row, ("v_cmd_x", "v_cmd_y", "v_cmd_z"))
        if not math.isfinite(t) or v is None or prev_t is None or prev_v is None:
            result.append(0.0 if math.isfinite(t) and v is not None else math.nan)
        else:
            dt = t - prev_t
            if dt <= 1.0e-6:
                result.append(math.nan)
            else:
                dv = math.sqrt(sum((v[index] - prev_v[index]) ** 2 for index in range(3)))
                result.append(dv / dt / GRAVITY_MPS2)
        if math.isfinite(t) and v is not None:
            prev_t = t
            prev_v = v
    return result


def _guidance_load(rows: list[dict[str, str]]) -> list[float]:
    result: list[float] = []
    for row in rows:
        n_cmd = _float(row.get("n_cmd_g"))
        if math.isfinite(n_cmd):
            result.append(n_cmd)
            continue
        vec = _vector(row, ("g_eval_x", "g_eval_y", "g_eval_z"))
        if vec is None:
            result.append(math.nan)
        else:
            result.append(math.sqrt(sum(component * component for component in vec)) / GRAVITY_MPS2)
    return result


def _unit(vector: np.ndarray) -> np.ndarray | None:
    norm = float(np.linalg.norm(vector))
    if not np.isfinite(norm) or norm <= 1.0e-9:
        return None
    return vector / norm


def _angle_deg(left: np.ndarray | None, right: np.ndarray | None) -> float:
    if left is None or right is None:
        return math.nan
    if not np.all(np.isfinite(left)) or not np.all(np.isfinite(right)):
        return math.nan
    dot = float(np.clip(np.dot(left, right), -1.0, 1.0))
    return math.degrees(math.acos(dot))


def _np_vector(row: dict[str, str], keys: tuple[str, str, str]) -> np.ndarray | None:
    values = np.array([_float(row.get(key)) for key in keys], dtype=float)
    if not np.all(np.isfinite(values)):
        return None
    return values


def _finite_difference_vectors(times: list[float], positions: list[np.ndarray | None]) -> list[np.ndarray | None]:
    velocities: list[np.ndarray | None] = [None] * len(times)
    if len(times) < 2:
        return velocities
    for index in range(len(times)):
        if index == 0:
            left, right = 0, 1
        elif index == len(times) - 1:
            left, right = len(times) - 2, len(times) - 1
        else:
            left, right = index - 1, index + 1
        if positions[left] is None or positions[right] is None:
            continue
        dt = times[right] - times[left]
        if math.isfinite(dt) and dt > 1.0e-5:
            velocities[index] = (positions[right] - positions[left]) / dt
    return velocities


def _camera_shadow_metrics(samples: list[dict[str, str]]) -> list[dict[str, float]]:
    times = [_float(row.get("t")) for row in samples]
    camera_pos = [_np_vector(row, ("camera_world_x", "camera_world_y", "camera_world_z")) for row in samples]
    intruder_pos = [_np_vector(row, ("intruder_x", "intruder_y", "intruder_z")) for row in samples]
    camera_vel = _finite_difference_vectors(times, camera_pos)
    intruder_vel = _finite_difference_vectors(times, intruder_pos)
    metrics: list[dict[str, float]] = []
    for index, row in enumerate(samples):
        rel = None if camera_pos[index] is None or intruder_pos[index] is None else intruder_pos[index] - camera_pos[index]
        rel_vel = None if camera_vel[index] is None or intruder_vel[index] is None else intruder_vel[index] - camera_vel[index]
        truth_los = _unit(rel) if rel is not None else None
        visual_vector = _np_vector(row, ("lambda_x", "lambda_y", "lambda_z"))
        visual_los = _unit(visual_vector) if visual_vector is not None else None
        range_m = float(np.linalg.norm(rel)) if rel is not None else math.nan
        closing_speed = math.nan
        lambda_dot_norm = math.nan
        shadow_vc_g = math.nan
        shadow_vm_g = math.nan
        if rel is not None and rel_vel is not None and math.isfinite(range_m) and range_m > 1.0e-6:
            los = rel / range_m
            closing_speed = -float(np.dot(rel_vel, los))
            omega = np.cross(rel, rel_vel) / max(range_m * range_m, 1.0e-12)
            lambda_dot = np.cross(omega, los)
            lambda_dot_norm = float(np.linalg.norm(lambda_dot))
            if closing_speed > 1.0e-3:
                shadow_vc_g = 3.0 * closing_speed * lambda_dot_norm / GRAVITY_MPS2
            else:
                shadow_vc_g = 0.0
            vm = _float(row.get("speed_ratio"), 2.0) * _float(row.get("intruder_speed"), 5.0)
            shadow_vm_g = 3.0 * vm * lambda_dot_norm / GRAVITY_MPS2
        metrics.append(
            {
                "t": times[index],
                "range_camera_m": range_m,
                "shadow_vc_n_req_g": shadow_vc_g,
                "shadow_vm_n_req_g": shadow_vm_g,
                "visual_los_error_deg": _angle_deg(visual_los, truth_los),
            }
        )
    return metrics


def _load_summary(path: Path, label: str) -> list[CaseRow]:
    if not path.exists():
        return []
    rows: list[CaseRow] = []
    for item in _read_csv(path):
        csv_path = _resolve(item.get("csv_path", ""))
        meta_path = _resolve(item.get("meta_path", "")) if item.get("meta_path") else csv_path.with_name(f"{csv_path.stem}_meta.json")
        meta = _read_json(meta_path)
        derived = meta.get("derived", {}) if isinstance(meta, dict) else {}
        rows.append(
            CaseRow(
                label=label,
                case=item.get("case", ""),
                start_range_m=_float(item.get("start_horizontal_range_m")),
                hit=_bool(item.get("hit")),
                hit_t_s=_float(item.get("hit_t_s")),
                near_hit=_bool(item.get("near_hit")),
                near_hit_t_s=_float(item.get("near_hit_t_s")),
                near_hit_range_m=_float(item.get("near_hit_range_m")),
                min_range_m=_float(item.get("min_range_m")),
                final_range_m=_float(item.get("final_range_m")),
                frames=_int(item.get("frames")),
                detected_frames=_int(item.get("detected_frames")),
                valid_frames=_int(item.get("valid_frames")),
                avg_wall_fps=_float(item.get("avg_wall_fps")),
                avg_sim_sample_fps=_float(item.get("avg_sim_sample_fps")),
                avg_detector_fps=_float(derived.get("avg_detector_fps")),
                max_load_factor_fd_g=_float(item.get("max_load_factor_fd_g")),
                csv_path=csv_path,
                meta_path=meta_path,
            )
        )
    return rows


def load_rows(stamp: str) -> list[CaseRow]:
    rows: list[CaseRow] = []
    for label in ("TTC", "VM"):
        rows.extend(_load_summary(LOG_DIR / f"yolo_sitl_{label}_{stamp}_summary.csv", label))
    return sorted(rows, key=lambda row: (row.start_range_m, row.label))


def _first_sample(rows: list[CaseRow]) -> dict[str, str]:
    for row in rows:
        if row.csv_path.exists():
            data = _read_csv(row.csv_path)
            if data:
                return data[0]
    return {}


def _first_meta(rows: list[CaseRow]) -> dict:
    for row in rows:
        meta = _read_json(row.meta_path)
        if meta:
            return meta
    return {}


def _summary_table(rows: list[CaseRow]) -> str:
    lines = [
        "|组别|碰撞命中|近距命中|碰撞距离m|近距距离m|未命中距离m|最小中心距离m|检测帧/总帧|有效帧/总帧|平均检测FPS|",
        "|---|---:|---:|---|---|---|---:|---:|---:|---:|",
    ]
    for label in ("TTC", "VM"):
        group = [row for row in rows if row.label == label]
        if not group:
            continue
        hit_ranges = ", ".join(f"{row.start_range_m:.0f}" for row in group if row.hit) or "-"
        near_hit_ranges = ", ".join(f"{row.start_range_m:.0f}" for row in group if row.near_hit) or "-"
        miss_ranges = ", ".join(f"{row.start_range_m:.0f}" for row in group if not row.hit and not row.near_hit) or "-"
        min_range = min([row.min_range_m for row in group if math.isfinite(row.min_range_m)] or [math.nan])
        frames = sum(row.frames for row in group)
        detected = sum(row.detected_frames for row in group)
        valid = sum(row.valid_frames for row in group)
        fps = [row.avg_detector_fps for row in group if math.isfinite(row.avg_detector_fps)]
        lines.append(
            f"|{label}|{sum(row.hit for row in group)}/{len(group)}|{sum(row.near_hit for row in group)}/{len(group)}|"
            f"{hit_ranges}|{near_hit_ranges}|{miss_ranges}|"
            f"{min_range:.3f}|{detected}/{frames}|{valid}/{frames}|"
            f"{(sum(fps) / len(fps) if fps else math.nan):.2f}|"
        )
    return "\n".join(lines)


def _detail_table(rows: list[CaseRow]) -> str:
    lines = [
        "|组别|距离m|碰撞|近距|碰撞时间s|近距时间s|近距距离m|最小距离m|终点距离m|检测帧率|有效帧率|YOLO FPS|sim FPS|实际过载max g|速度指令差分P95 g|需用过载P95 g|",
        "|---|---:|---:|---:|---:|---:|---:|---:|---:|---:|---:|---:|---:|---:|---:|---:|",
    ]
    for row in rows:
        samples = _read_csv(row.csv_path) if row.csv_path.exists() else []
        det_ratio = 100.0 * row.detected_frames / max(1, row.frames)
        valid_ratio = 100.0 * row.valid_frames / max(1, row.frames)
        lines.append(
            f"|{row.label}|{row.start_range_m:.0f}|{1 if row.hit else 0}|{1 if row.near_hit else 0}|"
            f"{'-' if not math.isfinite(row.hit_t_s) else f'{row.hit_t_s:.2f}'}|"
            f"{'-' if not math.isfinite(row.near_hit_t_s) else f'{row.near_hit_t_s:.2f}'}|"
            f"{'-' if not math.isfinite(row.near_hit_range_m) else f'{row.near_hit_range_m:.3f}'}|"
            f"{row.min_range_m:.3f}|{row.final_range_m:.3f}|{det_ratio:.1f}%|{valid_ratio:.1f}%|"
            f"{row.avg_detector_fps:.2f}|{row.avg_sim_sample_fps:.2f}|{row.max_load_factor_fd_g:.2f}|"
            f"{_percentile(_required_load(samples), 0.95):.2f}|{_percentile(_guidance_load(samples), 0.95):.2f}|"
        )
    return "\n".join(lines)


def _settings_markdown(rows: list[CaseRow], stamp: str) -> str:
    sample = _first_sample(rows)
    meta = _first_meta(rows).get("args", {})

    def cfg(key: str, default: object = "") -> object:
        value = sample.get(key, "")
        if value != "":
            return value
        return meta.get(key, default)

    return "\n".join(
        [
            "|参数|值|",
            "|---|---|",
            f"|stamp|`{stamp}`|",
            f"|settings|`{meta.get('settings_path', sample.get('settings_path', ''))}`|",
            f"|拦截机|`PX4 SITL / {cfg('px4_command_mode')}`|",
            f"|目标 actor|`{cfg('intruder_actor_name')}`|",
            f"|actor asset|`{cfg('intruder_actor_asset')}`|",
            f"|actor scale|`{cfg('intruder_actor_scale')}`|",
            f"|检测源|`{cfg('detector_source')}`|",
            f"|YOLO model|`{cfg('yolo_model')}`|",
            f"|YOLO device|`{cfg('yolo_device')}` runtime `{sample.get('yolo_runtime_device', '')}`|",
            f"|YOLO conf / iou / imgsz|`{cfg('yolo_conf')}` / `{cfg('yolo_iou')}` / `{cfg('yolo_imgsz')}`|",
            f"|tracker|`{cfg('yolo_tracker')}`，single target `{cfg('yolo_single_target_mode')}`|",
            f"|相机外参|`x={cfg('camera_x')}, y={cfg('camera_y')}, z={cfg('camera_z')}`|",
            f"|upward centering|`{cfg('upward_centering')}`, gain `{cfg('upward_centering_gain')}`, max accel `{cfg('upward_centering_max_accel_mps2')} m/s^2`|",
            f"|near-hit radius|`{cfg('near_hit_radius_m')} m`|",
            f"|FOV / resolution|`{cfg('fov_deg')} deg`, `{sample.get('image_width_runtime', sample.get('width', ''))}x{sample.get('image_height_runtime', sample.get('height', ''))}`|",
            f"|高度差|`{cfg('intruder_altitude_offset_m')} m`|",
            f"|目标速度 / speed ratio|`{cfg('intruder_speed')} m/s` / `{cfg('speed_ratio')}`|",
            f"|rate_hz|`{cfg('rate_hz')}`|",
            f"|guidance output|`{cfg('guidance_output_mode')}`|",
            f"|max guidance accel|`{cfg('max_guidance_accel_mps2')} m/s^2`|",
            f"|min speed ratio|`{cfg('min_speed_ratio')}`|",
            f"|thrust model|`{cfg('thrust_model')}`, mass `{cfg('vehicle_mass_kg')} kg`, max total thrust `{cfg('vehicle_max_total_thrust_n')} N`|",
            f"|body-rate tilt / attitude P|`{cfg('body_rate_max_tilt_deg')} deg` / `{cfg('body_rate_attitude_p')}`|",
            f"|body-rate roll/pitch max rate|`{cfg('body_rate_max_roll_rate_deg')}` / `{cfg('body_rate_max_pitch_rate_deg')} deg/s`|",
            f"|body-rate profile|`{cfg('body_rate_control_profile')}`|",
            f"|body-rate v2 Kp roll/pitch/yaw|`{cfg('body_rate_v2_kp_roll')}` / `{cfg('body_rate_v2_kp_pitch')}` / `{cfg('body_rate_v2_kp_yaw')}`|",
            f"|body-rate v2 max pq / slew pq-r|`{cfg('body_rate_v2_max_pq_rate_deg_s')} deg/s` / `{cfg('body_rate_v2_slew_pq_deg_s2')}`-`{cfg('body_rate_v2_slew_r_deg_s2')} deg/s^2`|",
            f"|body-rate v2 thrust reserve / guard|`{cfg('body_rate_v2_thrust_reserve')}` / error `{cfg('body_rate_v2_guard_error_ratio')}`, PNG scale `{cfg('body_rate_v2_guard_png_scale')}`, speed-hold scale `{cfg('body_rate_v2_guard_speed_hold_scale')}`|",
            f"|body-rate thrust|min/hover/max `{cfg('body_rate_min_thrust')}` / `{cfg('body_rate_hover_thrust')}` / `{cfg('body_rate_max_thrust')}`|",
            f"|body-rate speed hold|gain `{cfg('body_rate_speed_hold_gain')}`, max accel `{cfg('body_rate_speed_hold_max_accel_mps2')} m/s^2`, total limit `{cfg('body_rate_total_accel_limit_mps2')} m/s^2`|",
            f"|attitude tilt / yaw lookahead|`{cfg('attitude_max_tilt_deg')} deg` / `{cfg('attitude_yaw_lookahead_s')} s`|",
            f"|attitude thrust|min/hover/max `{cfg('attitude_min_thrust')}` / `{cfg('attitude_hover_thrust')}` / `{cfg('attitude_max_thrust')}`|",
            f"|attitude speed hold|gain `{cfg('attitude_speed_hold_gain')}`, max accel `{cfg('attitude_speed_hold_max_accel_mps2')} m/s^2`, total limit `{cfg('attitude_total_accel_limit_mps2')} m/s^2`|",
            f"|LOS filter|`{cfg('los_filter_enabled', meta.get('los_filter', ''))}`|",
            f"|LOS KF q lambda / lambda_dot|`{cfg('los_filter_process_lambda')}` / `{cfg('los_filter_process_lambda_dot')}`|",
            f"|LOS KF r / innovation gate|`{cfg('los_filter_measurement_noise')}` / `{cfg('los_filter_innovation_reject')}`|",
            f"|LOS terminal gate / delay|`{cfg('los_filter_terminal_innovation_reject')}` / `{cfg('los_delay_compensation_s')} s`|",
            f"|terminal image KF|predict `{cfg('terminal_image_kf_max_predict_s')} s`, reject `{cfg('terminal_image_kf_innovation_reject_rad')} rad`, soft reject `{cfg('terminal_image_kf_soft_reject_predict', False)}`|",
            f"|terminal image KF dynamics|accel noise `{cfg('terminal_image_kf_accel_noise_rad_s2')} rad/s^2`, max rate `{cfg('terminal_image_kf_max_rate_rad_s')} rad/s`|",
            f"|terminal velocity blind-push|`{cfg('terminal_velocity_blind_push')}`|",
            f"|terminal blind requires visual loss|`{cfg('terminal_blind_requires_visual_loss', False)}`|",
            f"|terminal accel hold|`{cfg('terminal_accel_hold')}`, window `{cfg('terminal_accel_hold_window_s')} s`, decay `{cfg('terminal_accel_decay_tau_s')} s`, max `{cfg('terminal_accel_hold_max_mps2')} m/s^2`|",
            f"|frame_guard|`{meta.get('frame_guard', '')}`|",
            f"|bbox noise|`{cfg('bbox_noise_enabled', meta.get('bbox_noise', ''))}`|",
        ]
    )


def plot_summary(rows: list[CaseRow], output: Path) -> None:
    fig, axes = plt.subplots(2, 2, figsize=(13, 8))
    ax_min, ax_hit, ax_detect, ax_fps = axes.flat
    for label in ("TTC", "VM"):
        group = [row for row in rows if row.label == label]
        if not group:
            continue
        x = [row.start_range_m for row in group]
        ax_min.plot(x, [row.min_range_m for row in group], marker="o", label=LABELS[label])
        ax_hit.plot(x, [1 if row.hit else 0 for row in group], marker="o", label=f"{LABELS[label]} collision")
        ax_hit.plot(x, [1 if row.near_hit else 0 for row in group], marker="x", linestyle="--", label=f"{LABELS[label]} near-hit")
        ax_detect.plot(x, [100.0 * row.detected_frames / max(1, row.frames) for row in group], marker="o", label=LABELS[label])
        ax_fps.plot(x, [row.avg_detector_fps for row in group], marker="o", label=LABELS[label])
    ax_min.set_title("Minimum true center range")
    ax_min.set_ylabel("m")
    ax_hit.set_title("Collision and near-hit")
    ax_hit.set_ylabel("success")
    ax_detect.set_title("YOLO detection frame ratio")
    ax_detect.set_ylabel("%")
    ax_fps.set_title("Detector FPS")
    ax_fps.set_ylabel("FPS")
    for ax in axes.flat:
        ax.set_xlabel("Initial horizontal range / m")
        ax.grid(True, alpha=0.3)
        ax.legend(fontsize=8)
    fig.tight_layout()
    fig.savefig(output, dpi=170)
    plt.close(fig)


def plot_per_distance(rows: list[CaseRow], output_dir: Path) -> dict[float, Path]:
    images: dict[float, Path] = {}
    for range_m in sorted({row.start_range_m for row in rows if math.isfinite(row.start_range_m)}):
        group = [row for row in rows if row.start_range_m == range_m]
        fig, axes = plt.subplots(5, 1, figsize=(12, 14), sharex=False)
        ax_range, ax_bbox, ax_ttc, ax_load, ax_fps = axes
        for row in group:
            samples = _read_csv(row.csv_path) if row.csv_path.exists() else []
            t = _series(samples, "t")
            label = LABELS[row.label]
            outcome = "collision" if row.hit else ("near-hit" if row.near_hit else "miss")
            ax_range.plot(t, _series(samples, "range"), linewidth=1.1, label=f"{label} {outcome}")
            ax_bbox.plot(t, _series(samples, "bbox_area"), linewidth=1.0, label=label)
            ax_ttc.plot(t, _series(samples, "ttc"), linewidth=1.0, label=label)
            ax_load.plot(t, _series(samples, "load_factor_fd_g"), linewidth=1.0, label=f"{label} actual")
            ax_load.plot(t, _guidance_load(samples), linewidth=0.9, linestyle="--", label=f"{label} required")
            ax_fps.plot(t, _series(samples, "detector_fps"), linewidth=1.0, label=label)
        ax_range.set_title(f"{range_m:.0f}m true center range")
        ax_range.set_ylabel("m")
        ax_bbox.set_title("BBox area ratio")
        ax_bbox.set_ylabel("area")
        ax_ttc.set_title("TTC estimate")
        ax_ttc.set_ylabel("s")
        ax_load.set_title("Actual overload and required overload")
        ax_load.set_ylabel("g")
        ax_fps.set_title("Detector FPS")
        ax_fps.set_xlabel("Time / s")
        ax_fps.set_ylabel("FPS")
        for ax in axes:
            ax.grid(True, alpha=0.3)
            ax.legend(fontsize=6, ncol=2)
        fig.tight_layout()
        path = output_dir / f"yolo_sitl_ttc_vm_{int(round(range_m)):03d}m.png"
        fig.savefig(path, dpi=170)
        plt.close(fig)
        images[range_m] = path
    return images


def plot_shadow_summary(rows: list[CaseRow], output: Path) -> None:
    fig, axes = plt.subplots(2, 2, figsize=(13, 8))
    ax_los, ax_shadow, ax_det, ax_near = axes.flat
    for label in ("TTC", "VM"):
        group = [row for row in rows if row.label == label]
        if not group:
            continue
        x_values: list[float] = []
        los_p95: list[float] = []
        shadow_vc_p95: list[float] = []
        shadow_vm_p95: list[float] = []
        near_det: list[float] = []
        near_no_det: list[float] = []
        for row in group:
            samples = _read_csv(row.csv_path) if row.csv_path.exists() else []
            if not samples:
                continue
            shadow = _camera_shadow_metrics(samples)
            nearest = min(samples, key=lambda sample: _float(sample.get("range"), 1.0e9))
            nearest_t = _float(nearest.get("t"))
            near_indices = [
                index
                for index, sample in enumerate(samples)
                if math.isfinite(nearest_t) and abs(_float(sample.get("t")) - nearest_t) <= 1.0
            ]
            x_values.append(row.start_range_m)
            los_p95.append(_percentile([item["visual_los_error_deg"] for item in shadow], 0.95))
            shadow_vc_p95.append(_percentile([item["shadow_vc_n_req_g"] for item in shadow], 0.95))
            shadow_vm_p95.append(_percentile([item["shadow_vm_n_req_g"] for item in shadow], 0.95))
            if near_indices:
                detected = sum(1 for index in near_indices if _float(samples[index].get("detected"), 0.0) > 0.0)
                near_det.append(100.0 * detected / len(near_indices))
                near_no_det.append(float(len(near_indices) - detected))
            else:
                near_det.append(math.nan)
                near_no_det.append(math.nan)
        ax_los.plot(x_values, los_p95, marker="o", label=LABELS[label])
        ax_shadow.plot(x_values, shadow_vc_p95, marker="o", label=f"{label} N*Vc")
        ax_shadow.plot(x_values, shadow_vm_p95, marker="x", linestyle="--", label=f"{label} N*Vm")
        ax_det.plot(x_values, near_det, marker="o", label=LABELS[label])
        ax_near.plot(x_values, near_no_det, marker="o", label=LABELS[label])
    ax_los.set_title("Visual LOS vs camera-truth LOS P95")
    ax_los.set_ylabel("deg")
    ax_shadow.set_title("Camera-truth shadow PNG required overload P95")
    ax_shadow.set_ylabel("g")
    ax_det.set_title("Detection ratio near closest approach")
    ax_det.set_ylabel("%")
    ax_near.set_title("No-detection frames near closest approach")
    ax_near.set_ylabel("frames in +/-1s")
    for ax in axes.flat:
        ax.set_xlabel("Initial horizontal range / m")
        ax.grid(True, alpha=0.3)
        ax.legend(fontsize=8)
    fig.tight_layout()
    fig.savefig(output, dpi=170)
    plt.close(fig)


def _shadow_diagnostics_markdown(rows: list[CaseRow]) -> str:
    lines = [
        "影子测试不参与导引，只用日志中的相机光心 `camera_world_*` 与目标真值位置离线计算经典 `N*Vc` 和固定 `N*Vm` PNG 理论需用过载，并和视觉 LOS、检测连续性对齐。",
        "",
        "|组别|距离m|碰撞|最小距离m|最近点检测率|最近点无检测帧|视觉LOS误差P95|影子N*Vc P95 g|影子N*Vm P95 g|视觉需用P95 g|实际过载max g|",
        "|---|---:|---:|---:|---:|---:|---:|---:|---:|---:|---:|",
    ]
    for row in sorted(rows, key=lambda item: (item.start_range_m, item.label)):
        samples = _read_csv(row.csv_path) if row.csv_path.exists() else []
        if not samples:
            continue
        shadow = _camera_shadow_metrics(samples)
        nearest = min(samples, key=lambda sample: _float(sample.get("range"), 1.0e9))
        nearest_t = _float(nearest.get("t"))
        near_indices = [
            index
            for index, sample in enumerate(samples)
            if math.isfinite(nearest_t) and abs(_float(sample.get("t")) - nearest_t) <= 1.0
        ]
        if near_indices:
            detected = sum(1 for index in near_indices if _float(samples[index].get("detected"), 0.0) > 0.0)
            near_detect_ratio = 100.0 * detected / len(near_indices)
            near_no_detection = len(near_indices) - detected
            near_los_p95 = _percentile([shadow[index]["visual_los_error_deg"] for index in near_indices], 0.95)
        else:
            near_detect_ratio = math.nan
            near_no_detection = 0
            near_los_p95 = math.nan
        lines.append(
            f"|{row.label}|{row.start_range_m:.0f}|{1 if row.hit else 0}|{row.min_range_m:.3f}|"
            f"{near_detect_ratio:.1f}%|{near_no_detection}/{len(near_indices)}|{near_los_p95:.1f}|"
            f"{_percentile([item['shadow_vc_n_req_g'] for item in shadow], 0.95):.2f}|"
            f"{_percentile([item['shadow_vm_n_req_g'] for item in shadow], 0.95):.2f}|"
            f"{_percentile(_guidance_load(samples), 0.95):.2f}|{row.max_load_factor_fd_g:.2f}|"
        )
    lines.extend(
        [
            "",
            "- 如果影子 `N*Vc` P95 很低但视觉 LOS 误差和无检测帧较高，优先定位检测连续性、LOS KF/外推和 frame-centering。",
            "- 如果视觉需用过载高而实际过载低，优先定位 PX4 姿态/推力响应、倾角限制和 speed-hold 混合项。",
        ]
    )
    return "\n".join(lines)


def _case_notes(rows: list[CaseRow]) -> str:
    lines: list[str] = []
    for label in ("TTC", "VM"):
        group = [row for row in rows if row.label == label]
        if not group:
            continue
        hit_ranges = ", ".join(f"{row.start_range_m:.0f}m" for row in group if row.hit) or "-"
        miss_ranges = ", ".join(f"{row.start_range_m:.0f}m" for row in group if not row.hit) or "-"
        det = sum(row.detected_frames for row in group) / max(1, sum(row.frames for row in group))
        valid = sum(row.valid_frames for row in group) / max(1, sum(row.frames for row in group))
        fps = [row.avg_detector_fps for row in group if math.isfinite(row.avg_detector_fps)]
        lines.append(
            f"- {label}: 命中 `{sum(row.hit for row in group)}/{len(group)}`，命中距离 `{hit_ranges}`，"
            f"未命中 `{miss_ranges}`，检测帧比例 `{100.0 * det:.1f}%`，有效导引帧比例 `{100.0 * valid:.1f}%`，"
            f"平均检测 FPS `{(sum(fps) / len(fps) if fps else math.nan):.2f}`。"
        )
    lines.append(
        "- 本轮使用真实 YOLOv8 + ByteTrack，因此检测连续性和 GPU 推理速度会直接进入闭环；结果不能和 AirSim detect 函数的理想 bbox 直接等价比较。"
    )
    lines.append(
        "- `accel_integral` 模式的 `n_cmd_g` 来自导引层 `a_cmd`，底层仍通过 PX4/AirSim 速度 setpoint 闭环；实际过载由真实速度差分估计，因此会同时受 PX4 响应、速度限幅和视觉帧率影响。"
    )
    lines.append(
        "- `accel_body_rate` 模式下 `n_cmd_g` 仍表示纯 PNG 需用过载；实际发送给 PX4 的是 `SET_ATTITUDE_TARGET` 机体系 `p/q/r` 角速度和归一化 thrust，日志中的 `body_rate_control_accel_*` 额外包含沿 LOS 的速度保持加速度。"
    )
    lines.append(
        "- `accel_attitude` 模式下 `n_cmd_g` 同样表示纯 PNG 需用过载；实际发送给 PX4 的是 `SET_ATTITUDE_TARGET` 姿态四元数和归一化 thrust，日志中的 `attitude_control_accel_*` 记录姿态指令生成前的合成加速度。"
    )
    return "\n".join(lines)


def _reason(row: dict[str, str]) -> str:
    return row.get("reject_reason") or row.get("reason") or "valid"


def _diagnostics_markdown(rows: list[CaseRow]) -> str:
    lines = [
        "|组别|距离m|最近距离m|最近点状态|主要失败/降级原因|检测率|有效率|",
        "|---|---:|---:|---|---|---:|---:|",
    ]
    for row in sorted(rows, key=lambda item: (item.start_range_m, item.label)):
        samples = _read_csv(row.csv_path) if row.csv_path.exists() else []
        if not samples:
            lines.append(f"|{row.label}|{row.start_range_m:.0f}|-|-|无 CSV|-|-|")
            continue
        nearest = min(samples, key=lambda sample: _float(sample.get("range"), 1.0e9))
        counts: dict[str, int] = {}
        for sample in samples:
            reason = _reason(sample)
            counts[reason] = counts.get(reason, 0) + 1
        common = ", ".join(f"{reason}:{count}" for reason, count in sorted(counts.items(), key=lambda item: item[1], reverse=True)[:4])
        det_ratio = 100.0 * row.detected_frames / max(1, row.frames)
        valid_ratio = 100.0 * row.valid_frames / max(1, row.frames)
        lines.append(
            f"|{row.label}|{row.start_range_m:.0f}|{_float(nearest.get('range')):.3f}|"
            f"`{_reason(nearest)}`|{common}|{det_ratio:.1f}%|{valid_ratio:.1f}%|"
        )
    sample = _first_sample(rows)
    near_misses = [row for row in rows if not row.hit and math.isfinite(row.min_range_m) and row.min_range_m <= 3.0]
    low_detection = [row for row in rows if row.frames > 0 and row.detected_frames / max(1, row.frames) < 0.6]
    nearest_rejects: list[str] = []
    actual_peak: list[float] = []
    required_p95: list[float] = []
    for row in rows:
        samples = _read_csv(row.csv_path) if row.csv_path.exists() else []
        if not samples:
            continue
        nearest = min(samples, key=lambda sample_row: _float(sample_row.get("range"), 1.0e9))
        reason = _reason(nearest)
        if not row.hit and reason not in {"valid", ""}:
            nearest_rejects.append(f"{row.label} {row.start_range_m:.0f}m:`{reason}`")
        actual_peak.append(row.max_load_factor_fd_g)
        required_p95.append(_percentile(_guidance_load(samples), 0.95))

    dynamic_lines = [
        "",
        (
            "- LOS KF 参数："
            f"`q_lambda={sample.get('los_filter_process_lambda', '')}`、"
            f"`q_lambda_dot={sample.get('los_filter_process_lambda_dot', '')}`、"
            f"`r={sample.get('los_filter_measurement_noise', '')}`、"
            f"`innovation_reject={sample.get('los_filter_innovation_reject', '')}`、"
            f"`terminal_reject={sample.get('los_filter_terminal_innovation_reject', '')}`。"
        ),
    ]
    if near_misses:
        dynamic_lines.append(
            "- 未命中但最近距离小于等于 3m 的工况："
            + "，".join(f"{row.label} {row.start_range_m:.0f}m({row.min_range_m:.3f}m)" for row in near_misses)
            + "。这些工况已接近目标，但没有触发 AirSim 碰撞判定，后续应重点看末端视场保持、外推和碰撞几何。"
        )
    if low_detection:
        dynamic_lines.append(
            "- 检测率低于 60% 的工况："
            + "，".join(
                f"{row.label} {row.start_range_m:.0f}m({100.0 * row.detected_frames / max(1, row.frames):.1f}%)"
                for row in low_detection
            )
            + "。这类失败优先归因于 YOLO/ByteTrack 连续性和固定相机视场保持，而不是导引律公式本身。"
        )
    if nearest_rejects:
        dynamic_lines.append(
            "- 最近点处仍处于降级或无效状态的未命中工况："
            + "，".join(nearest_rejects)
            + "。这些样本说明末端质量门、视觉外推和 bbox 裁切处理仍会影响命中窗口。"
        )
    if actual_peak and required_p95:
        actual_mean = sum(value for value in actual_peak if math.isfinite(value)) / max(1, len([value for value in actual_peak if math.isfinite(value)]))
        required_mean = sum(value for value in required_p95 if math.isfinite(value)) / max(1, len([value for value in required_p95 if math.isfinite(value)]))
        dynamic_lines.append(
            f"- 本轮平均实际过载峰值约 `{actual_mean:.2f} g`，平均需用过载 P95 约 `{required_mean:.2f} g`。"
            "两者不是同一个量：`n_cmd_g` 是导引层需求，实际过载还受 PX4 姿态/推力限制、YOLO 约 9 FPS 采样和 frame centering 限速影响。"
        )
    lines.extend(dynamic_lines)
    return "\n".join(lines)


def _body_rate_diagnostics_markdown(rows: list[CaseRow]) -> str:
    sample = _first_sample(rows)
    if sample.get("guidance_output_mode") != "accel_body_rate":
        return "本轮不是 `accel_body_rate` 输出模式，不生成 body-rate 控制诊断。"

    profile = str(sample.get("body_rate_control_profile") or "legacy")

    def row_values(samples: list[dict[str, str]], key: str) -> list[float]:
        return _finite(_series(samples, key))

    def active_ratio(samples: list[dict[str, str]], key: str) -> float:
        vals = row_values(samples, key)
        if not vals:
            return math.nan
        return 100.0 * sum(1 for value in vals if abs(value) > 1.0e-9) / len(vals)

    def peak_abs(samples: list[dict[str, str]], key: str) -> float:
        vals = row_values(samples, key)
        return max((abs(value) for value in vals), default=math.nan)

    if profile == "v2":
        lines = [
            "|组别|距离m|最近距离m|guard激活|p/q/r斜率限制|推力饱和|p/q/r峰值deg/s|姿态误差峰值|",
            "|---|---:|---:|---:|---|---:|---|---|",
        ]
        for row in sorted(rows, key=lambda item: (item.start_range_m, item.label)):
            samples = _read_csv(row.csv_path) if row.csv_path.exists() else []
            if not samples:
                continue

            guard = active_ratio(samples, "body_rate_frame_guard_active")
            slew_p = active_ratio(samples, "body_rate_p_slew_limited")
            slew_q = active_ratio(samples, "body_rate_q_slew_limited")
            slew_r = active_ratio(samples, "body_rate_r_slew_limited")
            thrust_sat = active_ratio(samples, "body_rate_thrust_saturated")
            p_peak = peak_abs(samples, "body_rate_p_deg_s")
            q_peak = peak_abs(samples, "body_rate_q_deg_s")
            r_peak = peak_abs(samples, "body_rate_r_deg_s")
            ex_peak = peak_abs(samples, "body_rate_q_error_x")
            ey_peak = peak_abs(samples, "body_rate_q_error_y")
            ez_peak = peak_abs(samples, "body_rate_q_error_z")
            lines.append(
                f"|{row.label}|{row.start_range_m:.0f}|{row.min_range_m:.3f}|{guard:.1f}%|"
                f"{slew_p:.1f}%/{slew_q:.1f}%/{slew_r:.1f}%|{thrust_sat:.1f}%|"
                f"{p_peak:.1f}/{q_peak:.1f}/{r_peak:.1f}|{ex_peak:.3f}/{ey_peak:.3f}/{ez_peak:.3f}|"
            )

        lines.extend(
            [
                "",
                "- `guard激活` 为 body-rate v2 的视场保持保护状态。该状态下 PNG 加速度和 speed-hold 加速度会分别乘以 `body_rate_v2_guard_png_scale` 与 `body_rate_v2_guard_speed_hold_scale`。",
                "- `p/q/r斜率限制` 表示 slew rate limit 介入比例；比例高时末端响应主要受角速度变化率约束。",
                "- `p/q/r峰值` 接近限幅且推力饱和不高时，应优先调角速度限幅、slew 和视觉闭环延迟，而不是只提高 PNG 增益。",
            ]
        )
        return "\n".join(lines)

    lines = [
        "|组别|距离m|最近距离m|frame-centering激活|推力饱和|p/q/r峰值deg/s|roll/pitch指令峰值deg|thrust min/max|",
        "|---|---:|---:|---:|---:|---|---|---|",
    ]
    for row in sorted(rows, key=lambda item: (item.start_range_m, item.label)):
        samples = _read_csv(row.csv_path) if row.csv_path.exists() else []
        if not samples:
            continue

        frame_centering = active_ratio(samples, "frame_centering_active")
        thrust_sat = active_ratio(samples, "body_rate_thrust_saturated")
        p_peak = peak_abs(samples, "body_rate_p_deg_s")
        q_peak = peak_abs(samples, "body_rate_q_deg_s")
        r_peak = peak_abs(samples, "body_rate_r_deg_s")
        roll_sp_peak = peak_abs(samples, "roll_sp_deg")
        pitch_sp_peak = peak_abs(samples, "pitch_sp_deg")
        thrust_vals = row_values(samples, "body_rate_thrust")
        thrust_min = min(thrust_vals, default=math.nan)
        thrust_max = max(thrust_vals, default=math.nan)
        lines.append(
            f"|{row.label}|{row.start_range_m:.0f}|{row.min_range_m:.3f}|{frame_centering:.1f}%|"
            f"{thrust_sat:.1f}%|{p_peak:.1f}/{q_peak:.1f}/{r_peak:.1f}|"
            f"{roll_sp_peak:.1f}/{pitch_sp_peak:.1f}|{thrust_min:.2f}/{thrust_max:.2f}|"
        )

    lines.extend(
        [
            "",
            f"- 本轮 `body_rate_control_profile={profile}`。legacy body-rate 使用欧拉角误差比例环，没有 v2 的 frame guard 降权和 slew rate limit。",
            "- `frame-centering激活` 表示固定上视相机进入视场保持/末端捕获/丢失保持状态的比例；比例高时，命中结果更受视场保持策略和 yaw 丢失外推影响。",
            "- `推力饱和` 与 `thrust min/max` 用于判断末端是否被垂向机动和总推力限制约束；若饱和高，应优先放宽 thrust、tilt 或垂向速度上限。",
        ]
    )
    return "\n".join(lines)


def _png_control_flow_markdown(rows: list[CaseRow]) -> str:
    sample = _first_sample(rows)
    guidance_output = sample.get("guidance_output_mode", "")
    px4_mode = sample.get("px4_command_mode", "")
    body_profile = sample.get("body_rate_control_profile", "")
    if guidance_output == "accel_body_rate":
        intro = (
            f"本轮实际使用 `guidance_output={guidance_output}`、`px4_command_mode={px4_mode}`、"
            f"`body_rate_control_profile={body_profile}`。因此控制链路是“PNG 需用加速度 -> 合成控制加速度 -> "
            "PX4 `SET_ATTITUDE_TARGET` 机体系 `p/q/r` 角速度 + thrust”。"
        )
    elif guidance_output == "accel_attitude":
        intro = (
            f"本轮实际使用 `guidance_output={guidance_output}`、`px4_command_mode={px4_mode}`。"
            "因此控制链路是“PNG 需用加速度 -> 姿态四元数 + thrust”。"
        )
    elif guidance_output == "accel_integral":
        intro = (
            f"本轮实际使用 `guidance_output={guidance_output}`、`px4_command_mode={px4_mode}`。"
            "因此控制链路是“PNG 需用加速度 -> 速度积分 setpoint”。"
        )
    else:
        intro = (
            f"本轮实际使用 `guidance_output={guidance_output}`、`px4_command_mode={px4_mode}`。"
            "因此控制链路保留为速度指令/航向率指令形式。"
        )

    return f"""{intro}

### 10.1 视觉量到 6D LOS

YOLOv8 + ByteTrack 输出 `bbox center=(u,v)`、`bbox area`、`track_id` 和置信度。bbox 中心先通过相机内参转换成相机坐标系单位射线：

```text
x_n = (u - cx) / fx
y_n = (v - cy) / fy
lambda_C = normalize([x_n, y_n, 1])
```

再使用相机外参和机体姿态转到惯性系：

```text
lambda_I = normalize(R_IB * R_BC * lambda_C)
```

其中 `R_BC` 是相机到机体的固定安装旋转，`R_IB` 是机体到惯性系的姿态。LOS 角速度由相邻 LOS 差分并投影到垂直 LOS 的平面得到：

```text
lambda_dot = project_perpendicular((lambda_I[k] - lambda_I[k-1]) / dt, lambda_I[k])
omega_LOS = lambda_I x lambda_dot
```

启用 LOS KF 时，滤波器输出平滑后的 `lambda_I` 和 `omega_LOS`；末端允许更松的 innovation gate，避免目标仍在检测框内时 PNG 加速度被过早清零。

### 10.2 PNG 生成需用加速度和需用过载

两种导引的共同输出都是导引层需用加速度 `a_cmd`：

```text
a_cmd = guidance_gain * (omega_LOS x lambda_I)
a_cmd = clip_norm(a_cmd, max_guidance_accel_mps2)
n_cmd_g = ||a_cmd|| / g
```

`omega_LOS x lambda_I` 给出垂直于视线的修正方向；`n_cmd_g` 是导引层需用过载，只表示 PNG 希望产生的机动强度。它不等于无人机真实过载，真实过载还受 PX4 姿态控制、推力限制、速度保持项、视觉帧率和 frame centering 限速影响。

TTC 组使用 bbox 面积扩张估计 `TTC ~= A / A_dot`，当前只把 TTC 用作增益调度和末端触发；当 TTC 无效但 LOS 有效时，仍保留 LOS/V_m soft guidance。V_m 组不使用 TTC，直接采用固定：

```text
guidance_gain = N * V_m
V_m = speed_ratio * intruder_speed
```

### 10.3 与速度保持项合成

在 `accel_attitude` 和 `accel_body_rate` 中，PNG 横向修正不再积分成速度指令。速度只作为沿 LOS 的保速参考：

```text
v_ref = speed_cap * lambda_I
a_speed_hold = K_v * (v_ref - v_current)
a_control_I = clip_norm(a_cmd + a_speed_hold, total_accel_limit)
```

`a_cmd` 是纯 PNG 需用加速度；`a_speed_hold` 是工程闭环项，用于避免飞机速度掉到无法追击或过度横向漂移。报告中的 `n_cmd_g` 仍来自 `a_cmd`，而 `attitude_control_accel_*` / `body_rate_control_accel_*` 记录合成后的控制加速度。

### 10.4 加速度到姿态四元数链路

在 `accel_attitude + mavlink_attitude` 下，程序先由图像中心误差和 LOS 水平投影得到期望航向：

```text
yaw_sp = current_yaw + yaw_rate_cmd * attitude_yaw_lookahead_s
```

随后把惯性系合成加速度旋转到期望 yaw 对应的水平坐标系，得到 roll/pitch setpoint，并发送姿态四元数和 thrust。

### 10.5 加速度到机体系角速度链路

在 `accel_body_rate + mavlink_body_rate` 下，程序先把 `a_control_I` 转到机体系：

```text
a_control_B = R_BI * a_control_I
```

再得到期望 roll/pitch。legacy 用欧拉角误差比例环：

```text
p_cmd = K_att * (roll_sp  - roll)
q_cmd = K_att * (pitch_sp - pitch)
r_cmd = yaw_rate_cmd
```

body-rate v2 使用四元数误差：

```text
q_err = inverse(q_current) * q_desired
e_att = 2 * q_err.xyz
p_raw = Kp_roll  * e_att.x
q_raw = Kp_pitch * e_att.y
r_raw = Kp_yaw   * e_att.z + yaw_rate_cmd
```

随后做角速度限幅和斜率限制。v2 不增加一阶 LPF，只保留 slew rate limit，避免末端视觉闭环额外相位滞后。MAVLink 仍使用 `SET_ATTITUDE_TARGET`，但 `type_mask` 忽略姿态四元数，只让 PX4 接收 `body_roll_rate/body_pitch_rate/body_yaw_rate` 和 thrust。

v2 的 thrust 使用向量投影法，并通过 `body_rate_v2_thrust_reserve` 预留电机差速余量：

```text
z_B_I = R_IB * [0, 0, 1]^T
required_specific_force_I = [-a_control_I.x, -a_control_I.y, g - a_control_I.z]
thrust_raw = mass * dot(required_specific_force_I, z_B_I) / max_total_thrust
```

### 10.6 本报告中过载曲线的含义

- `需用过载 n_cmd_g`：由 PNG 的 `a_cmd` 直接换算，是导引层希望产生的过载。
- `实际过载 max g`：由拦截机真实速度差分估计，体现 PX4 和 AirSim 动力学真正实现出的机动。
- `速度指令差分 P95 g`：兼容旧速度输出模式的指标；在 `accel_body_rate` 和 `accel_attitude` 下主要作为参考，不代表直接发送给 PX4 的控制量。

因此，若 `n_cmd_g` 很平滑但实际过载不足，问题通常在姿态/推力响应、速度保持、限幅或视觉低帧率；若 `n_cmd_g` 本身突变，则应优先检查 LOS/KF、bbox 裁切、丢检外推和 frame guard 状态切换。"""


def _export_docx(docx_path: Path | None = None) -> Path | None:
    try:
        from docx import Document
        from docx.enum.section import WD_ORIENT
        from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
        from docx.shared import Cm, Pt
    except ImportError:
        print("python-docx not installed; skip docx export")
        return None

    output_path = docx_path or REPORT_PATH.with_suffix(".docx")
    text = REPORT_PATH.read_text(encoding="utf-8")
    doc = Document()
    section = doc.sections[0]
    section.orientation = WD_ORIENT.PORTRAIT
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(1.5)
    section.bottom_margin = Cm(1.5)
    section.left_margin = Cm(1.3)
    section.right_margin = Cm(1.3)

    for style_name in ("Normal", "Body Text"):
        if style_name in doc.styles:
            style = doc.styles[style_name]
            style.font.name = "Arial"
            style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
            style.font.size = Pt(10.5)
    for level in range(1, 5):
        style_name = f"Heading {level}"
        if style_name in doc.styles:
            style = doc.styles[style_name]
            style.font.name = "Arial"
            style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
            style.font.bold = True

    heading_re = re.compile(r"^(#{1,6})\s+(.*)$")
    image_re = re.compile(r"^!\[([^\]]*)\]\(([^)]+)\)\s*$")
    token_re = re.compile(r"(`[^`]+`|\*\*[^*]+\*\*)")

    def add_runs(paragraph, value: str) -> None:
        pos = 0
        for match in token_re.finditer(value):
            if match.start() > pos:
                paragraph.add_run(value[pos : match.start()])
            token = match.group(0)
            if token.startswith("`"):
                run = paragraph.add_run(token[1:-1])
                run.font.name = "Consolas"
                run._element.rPr.rFonts.set(qn("w:eastAsia"), "Consolas")
            else:
                run = paragraph.add_run(token[2:-2])
                run.bold = True
            pos = match.end()
        if pos < len(value):
            paragraph.add_run(value[pos:])

    def set_cell_text(cell, value: str, *, bold: bool = False) -> None:
        cell.text = ""
        paragraph = cell.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = paragraph.add_run(value.replace("`", ""))
        run.bold = bold
        run.font.size = Pt(8)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

    def shade(cell) -> None:
        tc_pr = cell._tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:fill"), "D9EAF7")
        tc_pr.append(shd)

    lines = text.splitlines()
    index = 0
    in_code = False
    code_lines: list[str] = []
    while index < len(lines):
        line = lines[index].rstrip()
        if line.startswith("```"):
            if in_code:
                paragraph = doc.add_paragraph()
                run = paragraph.add_run("\n".join(code_lines))
                run.font.name = "Consolas"
                run._element.rPr.rFonts.set(qn("w:eastAsia"), "Consolas")
                run.font.size = Pt(9)
                code_lines = []
                in_code = False
            else:
                in_code = True
            index += 1
            continue
        if in_code:
            code_lines.append(line)
            index += 1
            continue
        if not line:
            index += 1
            continue
        heading = heading_re.match(line)
        if heading:
            level = min(len(heading.group(1)), 4)
            doc.add_heading(heading.group(2), level=level)
            index += 1
            continue
        image = image_re.match(line)
        if image:
            label, rel = image.groups()
            path = REPORT_PATH.parent / rel
            if path.exists():
                paragraph = doc.add_paragraph()
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                paragraph.add_run().add_picture(str(path), width=Cm(18.0))
                if label:
                    caption = doc.add_paragraph(label)
                    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
            index += 1
            continue
        if line.startswith("|"):
            table_lines: list[str] = []
            while index < len(lines) and lines[index].strip().startswith("|"):
                table_lines.append(lines[index].strip())
                index += 1
            raw_rows = [[cell.strip() for cell in item.strip("|").split("|")] for item in table_lines]
            if len(raw_rows) >= 2:
                body_rows = [raw_rows[0]] + raw_rows[2:]
                cols = max(len(row) for row in body_rows)
                table = doc.add_table(rows=len(body_rows), cols=cols)
                table.alignment = WD_TABLE_ALIGNMENT.CENTER
                table.style = "Table Grid"
                for row_i, row in enumerate(body_rows):
                    for col_i in range(cols):
                        cell = table.cell(row_i, col_i)
                        set_cell_text(cell, row[col_i] if col_i < len(row) else "", bold=row_i == 0)
                        if row_i == 0:
                            shade(cell)
                doc.add_paragraph()
            continue
        if line.startswith("- "):
            paragraph = doc.add_paragraph(style="List Bullet")
            add_runs(paragraph, line[2:])
            index += 1
            continue
        paragraph = doc.add_paragraph()
        add_runs(paragraph, line)
        index += 1

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)
    with ZipFile(output_path) as archive:
        media = [name for name in archive.namelist() if name.startswith("word/media/")]
    print(f"docx={output_path}")
    print(f"docx_media_count={len(media)}")
    return output_path


def write_report(
    rows: list[CaseRow],
    stamp: str,
    summary_img: Path,
    shadow_summary_img: Path,
    per_distance: dict[float, Path],
    *,
    title: str,
    range_note: str,
) -> None:
    REPORT_PATH.parent.mkdir(parents=True, exist_ok=True)
    image_lines = "\n".join(f"![{int(range_m)}m]({_rel(path)})" for range_m, path in sorted(per_distance.items()))
    report = f"""# {title}

## 1. 实验目的

按照此前已命中的 YOLO 案例配置，改用真正 PX4 SITL actor 场景，比较两种捷联视觉比例导引。本报告优先使用 `n_cmd_g` 作为需用过载；旧日志没有该字段时才回退到 `g_eval` 等效过载。

- `TTC` 组：`ttc_png`，TTC 只参与增益调度，并保留 LOS/Vm soft guidance。
- `VM` 组：`fixed_vm_png`，不使用 TTC，固定 `N * V_m` 导引增益。
- `accel_integral` 输出模式：导引律先计算 `a_cmd` / `n_cmd_g`，再按当前仿真步长积分为速度 setpoint；这不是直接向 PX4 发送加速度 setpoint。
- `accel_body_rate` 输出模式：导引律先计算 PNG 需用加速度，再转换为 PX4 `SET_ATTITUDE_TARGET` 机体系角速度 `p/q/r` 和 thrust；速度只作为沿 LOS 保速参考，不再把 PNG 横向修正直接加到速度指令上。
- `accel_attitude` 输出模式：导引律先计算 PNG 需用加速度，再转换为 PX4 `SET_ATTITUDE_TARGET` 姿态四元数和 thrust；速度只作为沿 LOS 保速参考。

{range_note}

## 2. 基准条件

{_settings_markdown(rows, stamp)}

## 3. 总览图

![summary]({_rel(summary_img)})

## 4. 汇总表

{_summary_table(rows)}

## 5. 明细表

{_detail_table(rows)}

## 6. 分距离曲线

每个距离一张图，包含真实中心距离、bbox 面积、TTC 估计、实际过载/需用过载和 YOLO 检测 FPS。

{image_lines}

## 7. LOS KF 与失败原因诊断

{_diagnostics_markdown(rows)}

## 8. 相机光心真值影子测试诊断

![shadow_summary]({_rel(shadow_summary_img)})

{_shadow_diagnostics_markdown(rows)}

## 9. body-rate 控制诊断

{_body_rate_diagnostics_markdown(rows)}

## 10. PNG 到过载、姿态和角速度的控制流程

{_png_control_flow_markdown(rows)}

## 11. 结论

{_case_notes(rows)}
"""
    REPORT_PATH.write_text(report, encoding="utf-8")


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description="Generate YOLO SITL TTC/Vm report.")
    parser.add_argument("--stamp", required=True)
    parser.add_argument("--report-path", default=str(REPORT_PATH))
    parser.add_argument("--asset-dir", default=str(ASSET_DIR))
    parser.add_argument("--title", default=DEFAULT_TITLE)
    parser.add_argument("--range-note", default=DEFAULT_RANGE_NOTE)
    parser.add_argument("--docx-path", default="")
    parser.add_argument("--no-docx", action="store_true")
    return parser.parse_args()


def main() -> None:
    global REPORT_PATH, ASSET_DIR
    args = parse_args()
    REPORT_PATH = _resolve(args.report_path)
    ASSET_DIR = _resolve(args.asset_dir)
    rows = load_rows(args.stamp)
    if not rows:
        raise SystemExit(f"no rows found for stamp {args.stamp}")
    ASSET_DIR.mkdir(parents=True, exist_ok=True)
    summary_img = ASSET_DIR / f"summary_{args.stamp}.png"
    shadow_summary_img = ASSET_DIR / f"shadow_summary_{args.stamp}.png"
    plot_summary(rows, summary_img)
    plot_shadow_summary(rows, shadow_summary_img)
    per_distance = plot_per_distance(rows, ASSET_DIR)
    write_report(rows, args.stamp, summary_img, shadow_summary_img, per_distance, title=args.title, range_note=args.range_note)
    print(f"report={REPORT_PATH}")
    print(f"stamp={args.stamp}")
    if not args.no_docx:
        _export_docx(_resolve(args.docx_path) if args.docx_path else REPORT_PATH.with_suffix(".docx"))


if __name__ == "__main__":
    main()
