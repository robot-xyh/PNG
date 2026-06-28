from __future__ import annotations

import argparse
import csv
import math
import re
import sys
from dataclasses import dataclass
from pathlib import Path

import matplotlib

matplotlib.use("Agg")
import matplotlib.pyplot as plt
import numpy as np


PROJECT_ROOT = Path(__file__).resolve().parents[1]
if str(PROJECT_ROOT) not in sys.path:
    sys.path.insert(0, str(PROJECT_ROOT))

GRAVITY_MPS2 = 9.80665
DEFAULT_TRAJECTORY_DIR = PROJECT_ROOT / "logs" / "body_rate_diagnostic"
DEFAULT_REPORT_PATH = PROJECT_ROOT / "完整方案" / "BodyRate_五组诊断实验报告.md"
DEFAULT_ASSET_DIR = PROJECT_ROOT / "完整方案" / "assets" / "BodyRate_五组诊断实验报告"

GROUP_LABELS = {
    "A_truth": "A 真值位置 PNG",
    "B_gimbal_detect": "B 云台相机 + AirSim detect",
    "C_strapdown_detect": "C 捷联相机 + AirSim detect",
    "D_strapdown_detect_high_authority": "D 捷联 detect + 高控制权限",
    "E_strapdown_detect_high_authority_20hz": "E 捷联 detect + 高控制权限 + 20Hz",
}


@dataclass
class Case:
    group: str
    law: str
    range_m: float
    csv_path: Path
    rows: list[dict[str, str]]
    hit: bool
    hit_t: float
    min_range: float
    final_range: float
    frames: int
    detected_rate: float
    valid_rate: float
    avg_wall_fps: float
    avg_sim_fps: float
    avg_detector_fps: float
    max_n_cmd_g: float
    max_actual_g: float
    body_rate_active_rate: float
    thrust_sat_rate: float


def _float(value: object, default: float = math.nan) -> float:
    if value is None or value == "":
        return default
    try:
        number = float(value)
    except (TypeError, ValueError):
        return default
    return number if math.isfinite(number) else default


def _read_csv(path: Path) -> list[dict[str, str]]:
    with path.open(newline="", encoding="utf-8") as stream:
        return list(csv.DictReader(stream))


def _finite(values: list[float]) -> list[float]:
    return [value for value in values if math.isfinite(value)]


def _mean(values: list[float]) -> float:
    finite = _finite(values)
    return float(sum(finite) / len(finite)) if finite else math.nan


def _max(values: list[float]) -> float:
    finite = _finite(values)
    return max(finite) if finite else math.nan


def _rate(rows: list[dict[str, str]], key: str) -> float:
    if not rows:
        return math.nan
    return sum(1 for row in rows if str(row.get(key, "")).strip() in {"1", "true", "True"}) / len(rows)


def _case_from_csv(path: Path, stamp: str) -> Case | None:
    pattern = re.compile(rf"^body_rate_diag_(?P<group>.+)_(?P<law>TTC|VM)_{re.escape(stamp)}_r(?P<range>[0-9.]+)_h")
    match = pattern.match(path.stem)
    if not match:
        return None
    rows = _read_csv(path)
    if not rows:
        return None
    ranges = [_float(row.get("range")) for row in rows]
    times = [_float(row.get("t")) for row in rows]
    hit_rows = [row for row in rows if str(row.get("hit", "")).strip() in {"1", "true", "True"}]
    hit = bool(hit_rows)
    hit_t = _float(hit_rows[0].get("t")) if hit_rows else math.nan
    min_range = min(_finite(ranges)) if _finite(ranges) else math.nan
    final_range = next((value for value in reversed(ranges) if math.isfinite(value)), math.nan)
    detector_fps = [_float(row.get("detector_fps")) for row in rows]
    if not _finite(detector_fps):
        detector_fps = [_float(row.get("yolo_fps")) for row in rows]
    return Case(
        group=match.group("group"),
        law=match.group("law"),
        range_m=float(match.group("range")),
        csv_path=path,
        rows=rows,
        hit=hit,
        hit_t=hit_t,
        min_range=min_range,
        final_range=final_range,
        frames=len(rows),
        detected_rate=_rate(rows, "detected"),
        valid_rate=_rate(rows, "valid"),
        avg_wall_fps=_mean([_float(row.get("wall_fps")) for row in rows]),
        avg_sim_fps=_mean([_float(row.get("sim_sample_fps")) for row in rows]),
        avg_detector_fps=_mean(detector_fps),
        max_n_cmd_g=_max([_float(row.get("n_cmd_g")) for row in rows]),
        max_actual_g=_max([_float(row.get("load_factor_fd_g")) for row in rows]),
        body_rate_active_rate=_rate(rows, "body_rate_control_active"),
        thrust_sat_rate=_rate(rows, "body_rate_thrust_saturated"),
    )


def _collect_cases(trajectory_dir: Path, stamp: str) -> list[Case]:
    cases = []
    for path in sorted(trajectory_dir.glob(f"body_rate_diag_*_{stamp}_r*.csv")):
        case = _case_from_csv(path, stamp)
        if case is not None:
            cases.append(case)
    return sorted(cases, key=lambda c: (c.group, c.law, c.range_m))


def _rel(path: Path, report_path: Path) -> str:
    try:
        return path.relative_to(report_path.parent).as_posix()
    except ValueError:
        return path.as_posix()


def _plot_summary(cases: list[Case], asset_dir: Path) -> dict[str, Path]:
    asset_dir.mkdir(parents=True, exist_ok=True)
    outputs: dict[str, Path] = {}

    fig, ax = plt.subplots(figsize=(11, 6))
    for key in sorted({(case.group, case.law) for case in cases}):
        subset = sorted([case for case in cases if (case.group, case.law) == key], key=lambda c: c.range_m)
        if not subset:
            continue
        ax.plot(
            [case.range_m for case in subset],
            [case.min_range for case in subset],
            marker="o",
            label=f"{key[0]} {key[1]}",
        )
    ax.axhline(1.0, color="tab:red", linestyle="--", linewidth=1.0, label="1m")
    ax.set_xlabel("Start horizontal range / m")
    ax.set_ylabel("Minimum true range / m")
    ax.set_title("Minimum Range by Group and Guidance Law")
    ax.grid(True, alpha=0.3)
    ax.legend(fontsize=8, ncol=2)
    fig.tight_layout()
    outputs["min_range"] = asset_dir / "min_range_summary.png"
    fig.savefig(outputs["min_range"], dpi=180)
    plt.close(fig)

    groups = sorted({case.group for case in cases})
    laws = sorted({case.law for case in cases})
    ranges = sorted({case.range_m for case in cases})
    fig, axes = plt.subplots(len(laws), 1, figsize=(12, max(4, 2.5 * len(laws))), squeeze=False)
    for row_index, law in enumerate(laws):
        ax = axes[row_index][0]
        matrix = np.full((len(groups), len(ranges)), np.nan)
        for case in cases:
            if case.law != law:
                continue
            matrix[groups.index(case.group), ranges.index(case.range_m)] = 1.0 if case.hit else 0.0
        image = ax.imshow(matrix, vmin=0.0, vmax=1.0, cmap="RdYlGn", aspect="auto")
        ax.set_xticks(range(len(ranges)), [f"{value:.0f}" for value in ranges])
        ax.set_yticks(range(len(groups)), [GROUP_LABELS.get(group, group) for group in groups])
        ax.set_title(f"Hit Matrix - {law}")
        for y in range(len(groups)):
            for x in range(len(ranges)):
                value = matrix[y, x]
                if math.isfinite(float(value)):
                    ax.text(x, y, "hit" if value >= 0.5 else "miss", ha="center", va="center", fontsize=8)
        fig.colorbar(image, ax=ax, fraction=0.025, pad=0.02)
    fig.tight_layout()
    outputs["hit_matrix"] = asset_dir / "hit_matrix.png"
    fig.savefig(outputs["hit_matrix"], dpi=180)
    plt.close(fig)

    for key in sorted({(case.group, case.law) for case in cases}):
        subset = sorted([case for case in cases if (case.group, case.law) == key], key=lambda c: c.range_m)
        fig, axes = plt.subplots(2, 1, figsize=(11, 7), sharex=False)
        for case in subset:
            times = [_float(row.get("t")) for row in case.rows]
            ranges = [_float(row.get("range")) for row in case.rows]
            axes[0].plot(times, ranges, label=f"{case.range_m:.0f}m {'hit' if case.hit else 'miss'}")
            n_cmd = [_float(row.get("n_cmd_g")) for row in case.rows]
            actual_g = [_float(row.get("load_factor_fd_g")) for row in case.rows]
            axes[1].plot(times, n_cmd, linewidth=1.0, label=f"{case.range_m:.0f}m cmd")
            axes[1].plot(times, actual_g, linewidth=0.8, linestyle="--", label=f"{case.range_m:.0f}m actual")
        axes[0].set_title(f"{GROUP_LABELS.get(key[0], key[0])} {key[1]} Range")
        axes[0].set_ylabel("Range / m")
        axes[0].grid(True, alpha=0.3)
        axes[1].set_title("Commanded vs Actual Load Factor")
        axes[1].set_xlabel("Time / s")
        axes[1].set_ylabel("g")
        axes[1].grid(True, alpha=0.3)
        axes[0].legend(fontsize=7, ncol=2)
        axes[1].legend(fontsize=7, ncol=2)
        fig.tight_layout()
        safe = f"{key[0]}_{key[1]}".replace("/", "_")
        outputs[safe] = asset_dir / f"{safe}_timeseries.png"
        fig.savefig(outputs[safe], dpi=180)
        plt.close(fig)
    return outputs


def _table(cases: list[Case]) -> list[str]:
    lines = [
        "|组别|导引|距离/m|是否碰撞|碰撞时刻/s|最小距离/m|最终距离/m|检测率|有效率|body-rate有效率|平均FPS|仿真采样FPS|最大指令过载/g|最大实际过载/g|推力饱和率|CSV|",
        "|---|---:|---:|---:|---:|---:|---:|---:|---:|---:|---:|---:|---:|---:|---:|---|",
    ]
    for case in cases:
        lines.append(
            "|"
            + "|".join(
                [
                    GROUP_LABELS.get(case.group, case.group),
                    case.law,
                    f"{case.range_m:.0f}",
                    "是" if case.hit else "否",
                    "" if not math.isfinite(case.hit_t) else f"{case.hit_t:.2f}",
                    "" if not math.isfinite(case.min_range) else f"{case.min_range:.2f}",
                    "" if not math.isfinite(case.final_range) else f"{case.final_range:.2f}",
                    "" if not math.isfinite(case.detected_rate) else f"{100*case.detected_rate:.1f}%",
                    "" if not math.isfinite(case.valid_rate) else f"{100*case.valid_rate:.1f}%",
                    "" if not math.isfinite(case.body_rate_active_rate) else f"{100*case.body_rate_active_rate:.1f}%",
                    "" if not math.isfinite(case.avg_wall_fps) else f"{case.avg_wall_fps:.1f}",
                    "" if not math.isfinite(case.avg_sim_fps) else f"{case.avg_sim_fps:.1f}",
                    "" if not math.isfinite(case.max_n_cmd_g) else f"{case.max_n_cmd_g:.2f}",
                    "" if not math.isfinite(case.max_actual_g) else f"{case.max_actual_g:.2f}",
                    "" if not math.isfinite(case.thrust_sat_rate) else f"{100*case.thrust_sat_rate:.1f}%",
                    f"`{case.csv_path.name}`",
                ]
            )
            + "|"
        )
    return lines


def _write_docx(report_path: Path, figures: dict[str, Path], cases: list[Case]) -> None:
    try:
        from docx import Document
        from docx.shared import Cm
    except Exception:
        print("python-docx not installed; skip docx export")
        return
    doc = Document()
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.left_margin = Cm(1.6)
    section.right_margin = Cm(1.6)
    section.top_margin = Cm(1.6)
    section.bottom_margin = Cm(1.6)
    doc.add_heading("Body-rate 五组诊断实验报告", 0)
    doc.add_paragraph(f"生成时间：{__import__('time').strftime('%Y-%m-%d %H:%M:%S')}")
    hit_count = sum(1 for case in cases if case.hit)
    doc.add_paragraph(f"总工况：{len(cases)}，碰撞成功：{hit_count}。")
    for name in ("hit_matrix", "min_range"):
        if name in figures:
            doc.add_picture(str(figures[name]), width=Cm(17.5))
    table = doc.add_table(rows=1, cols=7)
    hdr = table.rows[0].cells
    for index, title in enumerate(["组别", "导引", "距离", "命中", "最小距离", "检测率", "body-rate率"]):
        hdr[index].text = title
    for case in cases:
        cells = table.add_row().cells
        values = [
            GROUP_LABELS.get(case.group, case.group),
            case.law,
            f"{case.range_m:.0f}",
            "是" if case.hit else "否",
            "" if not math.isfinite(case.min_range) else f"{case.min_range:.2f}",
            "" if not math.isfinite(case.detected_rate) else f"{100*case.detected_rate:.1f}%",
            "" if not math.isfinite(case.body_rate_active_rate) else f"{100*case.body_rate_active_rate:.1f}%",
        ]
        for index, value in enumerate(values):
            cells[index].text = value
    for key, path in figures.items():
        if key in {"hit_matrix", "min_range"}:
            continue
        doc.add_heading(key, level=2)
        doc.add_picture(str(path), width=Cm(17.5))
    docx_path = report_path.with_suffix(".docx")
    doc.save(docx_path)
    print(f"docx={docx_path}")


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("--stamp", required=True)
    parser.add_argument("--trajectory-dir", default=str(DEFAULT_TRAJECTORY_DIR))
    parser.add_argument("--report-path", default=str(DEFAULT_REPORT_PATH))
    parser.add_argument("--asset-dir", default=str(DEFAULT_ASSET_DIR))
    parser.add_argument("--no-docx", action="store_true")
    args = parser.parse_args()

    trajectory_dir = Path(args.trajectory_dir).expanduser().resolve()
    report_path = Path(args.report_path).expanduser().resolve()
    asset_dir = Path(args.asset_dir).expanduser().resolve()
    cases = _collect_cases(trajectory_dir, args.stamp)
    if not cases:
        raise SystemExit(f"No body-rate diagnostic CSV files found for stamp={args.stamp} in {trajectory_dir}")
    figures = _plot_summary(cases, asset_dir)
    report_path.parent.mkdir(parents=True, exist_ok=True)
    hit_count = sum(1 for case in cases if case.hit)
    lines = [
        "# Body-rate 五组诊断实验报告",
        "",
        f"- 数据批次：`{args.stamp}`",
        f"- CSV 目录：`{trajectory_dir}`",
        f"- 总工况：{len(cases)}",
        f"- 碰撞成功：{hit_count}/{len(cases)}",
        "",
        "## 实验组定义",
        "",
        "- A：使用目标真实位置计算 PNG，只检验 body-rate 控制链路是否跑通。",
        "- B：云台相机持续锁定目标，AirSim detect 提供检测框，检验是否主要由捷联视场丢失造成。",
        "- C：捷联相机基线，AirSim detect 提供理论检测框，隔离 YOLO 识别断续影响。",
        "- D：在 C 基础上提高 PX4 速度、加速度、倾角与 body-rate 权限，检验控制权限瓶颈。",
        "- E：在 D 基础上提升控制循环到 20Hz，检验仿真/控制刷新率瓶颈。",
        "",
        "## 总览图",
        "",
        f"![hit_matrix]({_rel(figures['hit_matrix'], report_path)})",
        "",
        f"![min_range]({_rel(figures['min_range'], report_path)})",
        "",
        "## 结果表",
        "",
        *_table(cases),
        "",
        "## 分组曲线",
        "",
    ]
    for key, path in figures.items():
        if key in {"hit_matrix", "min_range"}:
            continue
        lines.extend([f"### `{key}`", "", f"![{key}]({_rel(path, report_path)})", ""])
    lines.extend(
        [
            "## 读数说明",
            "",
            "- `最大指令过载/g` 来自 PNG 加速度指令 `n_cmd_g`，是真值/视觉导引给控制层的需用过载。",
            "- `最大实际过载/g` 来自拦截机速度差分 `load_factor_fd_g`，反映 PX4/仿真实际响应。",
            "- `body-rate有效率` 小于 100% 通常表示该工况存在未进入控制、碰撞后截断、检测/导引无效或脚本提前退出。",
            "- B/C/D/E 使用 AirSim detect 时检测框是理论检测输出，若仍失败，优先检查视场保持、PX4 响应和 body-rate 控制映射。",
        ]
    )
    report_path.write_text("\n".join(lines) + "\n", encoding="utf-8")
    print(f"report={report_path}")
    if not args.no_docx:
        _write_docx(report_path, figures, cases)


if __name__ == "__main__":
    main()
